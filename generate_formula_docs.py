"""
Generate Formula Documentation Word files for:
1. Cost Factor Score Calc Formula
2. KPI Factor Score Calc Formula
3. Impact Value Score Calc Formula
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_DIR = r"C:\Users\krish\OneDrive\Desktop\impact-value"

# ─────────────────────────────────────────────────────────────────────────────
# Shared data
# ─────────────────────────────────────────────────────────────────────────────

BARRIER_TITLES = [
    "Lack of Training for Workers and Managers",
    "Resistance to Change",
    "Lack of Digital Culture and Training",
    "Higher Investment in Employees' Training",
    "Lack of Knowledge Management Systems",
    "Regulatory Compliance Issues",
    "Lack of Standards and Reference Architecture",
    "Lack of Internet Coverage and IT Facilities",
    "Limited Access to Funding and Credit",
    "Future Viability & Profitability",
    "Dependency on External Vendors",
    "High Implementation Cost",
    "Compliance with Sector-Specific Regulations",
    "Lack of Regulations and Standards",
    "Customers are Hesitant to Share Data",
]

COST_FACTORS = [
    ("1",  "Aftermarket Services / Warranty",               "aftermarket_services_warranty"),
    ("2",  "Depreciation",                                  "depreciation"),
    ("3",  "Labour",                                        "labour"),
    ("4",  "Maintenance & Repair",                          "maintenance_repair"),
    ("5",  "Raw Materials / Consumables",                   "raw_materials_consumables"),
    ("6",  "Rental / Operating Lease",                      "rental_operating_lease"),
    ("7",  "Research & Development",                        "research_development"),
    ("8",  "Selling, General & Administrative Expense",     "selling_general_administrative_expense"),
    ("9",  "Utilities",                                     "utilities"),
    ("10", "Earnings Before Interest & Taxes (EBIT)",       "earnings_before_interest_taxes_ebit"),
    ("11", "Financing Costs / Interest",                    "financing_costs_interest"),
    ("12", "Taxation & Compliance Costs",                   "taxation_compliance_costs"),
    ("13", "Supply Chain & Logistics Costs",                "supply_chain_logistics_costs"),
    ("14", "Technology & Digital Infrastructure Costs",     "technology_digital_infrastructure_costs"),
    ("15", "Training & Skill Development Costs",            "training_skill_development_costs"),
    ("16", "Regulatory Compliance Costs",                   "regulatory_compliance_costs"),
    ("17", "Insurance Costs",                               "insurance_costs"),
    ("18", "Marketing & Customer Acquisition Costs",        "marketing_customer_acquisition_costs"),
    ("19", "Environmental & Social Responsibility Costs",   "environmental_social_responsibility_costs"),
    ("20", "Quality Control & Assurance",                   "quality_control_assurance"),
]

KPI_FACTORS = [
    ("1",  "Asset & Equipment Efficiency",          "asset_equipment_efficiency"),
    ("2",  "Utilities Efficiency",                  "utilities_efficiency"),
    ("3",  "Inventory Efficiency",                  "inventory_efficiency"),
    ("4",  "Process Quality",                       "process_quality"),
    ("5",  "Product Quality",                       "product_quality"),
    ("6",  "Safety & Security",                     "safety_security"),
    ("7",  "Planning & Scheduling Effectiveness",   "planning_scheduling_effectiveness"),
    ("8",  "Time to Market",                        "time_to_market"),
    ("9",  "Production Flexibility",                "production_flexibility"),
    ("10", "Customer Satisfaction",                 "customer_satisfaction"),
    ("11", "Supply Chain Efficiency",               "supply_chain_efficiency"),
    ("12", "Market Share Growth",                   "market_share_growth"),
    ("13", "Employee Productivity",                 "employee_productivity"),
    ("14", "Return on Investment (ROI)",            "return_on_investment_roi"),
    ("15", "Financial Health & Stability",          "financial_health_and_stability"),
    ("16", "Talent Retention",                      "talent_retention"),
    ("17", "Customer Retention Rate",               "customer_retention_rate"),
]

# Weight matrices (from cost_service.py and kpi_service.py)
# Rows = Barriers 1-15,  Columns = cost/KPI factors in order above

COST_WEIGHTS = [
    [0, 0, 2, 0, 0, 0, 2, 1, 0, 2, 1, 1, 1, 3, 3, 1, 1, 1, 2, 1],  # B1
    [0, 0, 2, 0, 0, 0, 1, 1, 0, 1, 1, 1, 1, 2, 2, 1, 1, 1, 1, 1],  # B2
    [0, 0, 2, 0, 0, 0, 1, 1, 0, 1, 1, 1, 1, 3, 3, 1, 1, 1, 2, 1],  # B3
    [0, 0, 2, 0, 0, 0, 2, 1, 0, 1, 1, 1, 1, 3, 3, 1, 1, 1, 2, 1],  # B4
    [0, 0, 2, 0, 0, 0, 2, 1, 0, 2, 1, 1, 1, 3, 3, 1, 1, 1, 2, 1],  # B5
    [0, 0, 2, 0, 0, 0, 1, 1, 0, 1, 1, 1, 1, 2, 2, 1, 1, 1, 1, 1],  # B6
    [0, 0, 2, 0, 0, 0, 1, 1, 0, 1, 1, 1, 1, 2, 2, 1, 1, 1, 1, 1],  # B7
    [0, 0, 2, 0, 0, 0, 1, 1, 0, 1, 1, 1, 1, 2, 2, 1, 1, 1, 1, 1],  # B8
    [0, 0, 2, 0, 0, 0, 1, 2, 0, 3, 3, 1, 2, 1, 1, 1, 1, 1, 1, 1],  # B9
    [2, 1, 3, 2, 2, 1, 3, 3, 2, 3, 2, 2, 3, 3, 3, 3, 1, 3, 2, 3],  # B10
    [0, 0, 2, 0, 0, 0, 1, 1, 0, 2, 2, 1, 2, 1, 1, 1, 1, 1, 1, 1],  # B11
    [0, 0, 2, 0, 0, 0, 2, 1, 0, 3, 2, 3, 2, 1, 1, 3, 1, 1, 1, 1],  # B12
    [0, 0, 2, 0, 0, 0, 1, 1, 0, 2, 2, 3, 2, 1, 1, 3, 1, 1, 1, 1],  # B13
    [0, 0, 2, 0, 0, 0, 1, 1, 0, 2, 2, 3, 2, 1, 1, 3, 1, 1, 1, 1],  # B14
    [0, 0, 2, 0, 0, 0, 1, 1, 0, 2, 1, 1, 2, 3, 3, 1, 1, 1, 1, 1],  # B15
]

KPI_WEIGHTS = [
    [1, 1, 1, 2, 2, 1, 2, 1, 2, 1, 1, 2, 3, 2, 2, 3, 2],  # B1
    [1, 1, 1, 2, 2, 1, 2, 1, 2, 1, 1, 2, 3, 2, 2, 3, 2],  # B2
    [1, 1, 1, 2, 2, 1, 2, 2, 2, 1, 1, 1, 3, 2, 2, 3, 2],  # B3
    [1, 1, 1, 2, 2, 1, 2, 2, 2, 1, 1, 1, 3, 2, 2, 3, 2],  # B4
    [1, 1, 1, 2, 2, 1, 2, 2, 2, 1, 1, 1, 3, 2, 2, 2, 2],  # B5
    [1, 0, 1, 2, 2, 2, 1, 1, 2, 2, 1, 2, 1, 3, 3, 1, 2],  # B6
    [1, 0, 1, 2, 2, 2, 1, 1, 2, 2, 1, 2, 1, 2, 2, 1, 2],  # B7
    [1, 0, 1, 2, 2, 1, 2, 2, 2, 1, 1, 1, 2, 2, 2, 1, 1],  # B8
    [2, 1, 2, 1, 1, 0, 1, 2, 2, 1, 2, 2, 2, 3, 3, 1, 1],  # B9
    [3, 2, 3, 3, 3, 2, 3, 3, 3, 3, 3, 3, 3, 3, 3, 2, 3],  # B10
    [2, 1, 2, 1, 1, 1, 2, 2, 2, 1, 3, 1, 2, 2, 2, 1, 2],  # B11
    [1, 0, 1, 1, 1, 0, 1, 2, 1, 1, 2, 1, 1, 3, 3, 1, 1],  # B12
    [1, 0, 1, 2, 2, 2, 1, 1, 2, 2, 1, 2, 1, 3, 3, 1, 2],  # B13
    [1, 0, 1, 2, 2, 2, 1, 1, 2, 2, 1, 2, 1, 3, 3, 1, 2],  # B14
    [1, 0, 1, 2, 2, 2, 1, 1, 2, 2, 1, 1, 1, 2, 2, 1, 3],  # B15
]

# Impact Value weights (from settings.py)
WEIGHT_BARRIER  = 0.3
WEIGHT_COST     = 0.3
WEIGHT_KPI      = 0.4

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), '000000')
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # approx twips
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def header_row(table, texts, bg='1F4E79'):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(text)
        run.text = text
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, bg)
        set_cell_borders(cell)


def data_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9, italic=False, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if bg:
        set_cell_bg(cell, bg)
    set_cell_borders(cell)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    elif level == 1:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_formula_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Document 1: Cost Factor Score Calc Formula
# ─────────────────────────────────────────────────────────────────────────────

def make_cost_factor_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)

    # Title
    add_heading(doc, "Cost Factor Score Calculation Formula", level=0)
    doc.add_paragraph()

    # Overview
    add_heading(doc, "Overview", level=1)
    add_body(doc,
        "The Cost Factor Score measures how much each of the 20 cost factors "
        "contributes to the difficulty of overcoming each IoT adoption barrier. "
        "Users rate each cost factor on a scale of 1 (Very Low Impact) to 5 "
        "(Very High Impact). Each barrier has a pre-defined weight (0, 1, 2, or 3) "
        "for every cost factor. The higher the weight, the more that cost factor "
        "influences that specific barrier.")
    doc.add_paragraph()

    # Section 1 – Input Variables
    add_heading(doc, "Section 1: Input Variables (Cost Factors)", level=1)
    add_body(doc, "Each cost factor below is rated by the user on a scale of 1 to 5.")
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    header_row(tbl, ["#", "Cost Factor Name", "Variable Name"])
    tbl.rows[0].cells[0].width = Cm(1)
    for num, name, var in COST_FACTORS:
        row = tbl.add_row()
        data_cell(row.cells[0], num,  align=WD_ALIGN_PARAGRAPH.CENTER)
        data_cell(row.cells[1], name)
        data_cell(row.cells[2], var, italic=True)
    doc.add_paragraph()

    # Section 2 – Score Formula
    add_heading(doc, "Section 2: Score Formula", level=1)
    add_body(doc,
        "For each barrier i (where i = 1 to 15), the Cost Factor Score is calculated as "
        "the weighted sum of all 20 cost factor input values:")
    doc.add_paragraph()

    add_formula_para(doc,
        "Cost_Factor_Score(Barrier_i)  =  \u03A3 [CF_j  \u00D7  W(i, j)]   for j = 1 to 20")

    doc.add_paragraph()
    add_body(doc, "Where:")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("CF_j").italic = True
    r2 = p.add_run("  = Input value (1\u20135) of Cost Factor j")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("W(i, j)").italic = True
    r3 = p2.add_run("  = Weight of Cost Factor j for Barrier i  (values: 0, 1, 2, or 3)")
    r3.font.size = Pt(10)

    doc.add_paragraph()
    add_body(doc,
        "Weight Meaning:  0 = No relevance  |  1 = Low relevance  |  "
        "2 = Medium relevance  |  3 = High relevance")
    doc.add_paragraph()

    # Section 3 – Weight Matrix
    add_heading(doc, "Section 3: Barrier \u00D7 Cost Factor Weight Matrix", level=1)
    add_body(doc,
        "The table below shows the weight W(i, j) assigned to each cost factor "
        "for each barrier. Cells shaded in dark blue indicate a weight of 3 (High); "
        "medium blue indicates 2; light indicates 1; and blank (0) means not applicable.")
    doc.add_paragraph()

    # Build weight matrix table: rows = barriers, cols = cost factors
    num_cf = len(COST_FACTORS)
    tbl2 = doc.add_table(rows=1, cols=num_cf + 2)
    tbl2.style = 'Table Grid'

    # Header row
    hrow = tbl2.rows[0]
    data_cell(hrow.cells[0], "B#",    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg='1F4E79')
    hrow.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    data_cell(hrow.cells[1], "Barrier Title", bold=True, bg='1F4E79')
    hrow.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for j, (num, name, var) in enumerate(COST_FACTORS):
        cell = hrow.cells[j + 2]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(num)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, '1F4E79')
        set_cell_borders(cell)

    # Weight colours
    W_COLORS = {0: None, 1: 'D6E4F0', 2: '2E74B5', 3: '1F4E79'}
    W_TEXT_COLORS = {0: '000000', 1: '000000', 2: 'FFFFFF', 3: 'FFFFFF'}

    for i, (title, weights) in enumerate(zip(BARRIER_TITLES, COST_WEIGHTS)):
        row = tbl2.add_row()
        data_cell(row.cells[0], str(i + 1), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        data_cell(row.cells[1], title, size=8)
        for j, w in enumerate(weights):
            cell = row.cells[j + 2]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(w) if w > 0 else '')
            run.font.size = Pt(8)
            run.bold = (w == 3)
            bg = W_COLORS[w]
            if bg:
                set_cell_bg(cell, bg)
                run.font.color.rgb = RGBColor(
                    int(W_TEXT_COLORS[w][0:2], 16),
                    int(W_TEXT_COLORS[w][2:4], 16),
                    int(W_TEXT_COLORS[w][4:6], 16))
            set_cell_borders(cell)

    doc.add_paragraph()

    # Section 4 – Per Barrier Expanded Formula (matching original PDF table style)
    add_heading(doc, "Section 4: Per-Barrier Detailed Formula Table", level=1)
    add_body(doc,
        "The table below lists every cost factor with its weight for each barrier, "
        "following the same structure as the Barrier Score Calc Formula document.")
    doc.add_paragraph()

    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = 'Table Grid'
    header_row(tbl3, ["Barrier\nNumber", "Barrier Title", "Cost Factor", "Weight W(i,j)"])

    ROW_BG = ['F2F7FC', 'FFFFFF']
    for i, (title, weights) in enumerate(zip(BARRIER_TITLES, COST_WEIGHTS)):
        for j, w in enumerate(weights):
            row = tbl3.add_row()
            bg = ROW_BG[i % 2]
            data_cell(row.cells[0], str(i + 1), align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
            data_cell(row.cells[1], title, size=8, bg=bg)
            data_cell(row.cells[2], COST_FACTORS[j][1], size=8, bg=bg)
            w_text = str(w)
            cell_bg = bg if w < 3 else 'D6E4F0'
            data_cell(row.cells[3], w_text, align=WD_ALIGN_PARAGRAPH.CENTER,
                      bold=(w == 3), bg=cell_bg)

    path = f"{OUTPUT_DIR}\\Cost Factor Score Calc Formula.docx"
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# Document 2: KPI Factor Score Calc Formula
# ─────────────────────────────────────────────────────────────────────────────

def make_kpi_factor_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)

    add_heading(doc, "KPI Factor Score Calculation Formula", level=0)
    doc.add_paragraph()

    add_heading(doc, "Overview", level=1)
    add_body(doc,
        "The KPI Factor Score measures how strongly each Key Performance Indicator "
        "is affected by each IoT adoption barrier. Users rate each KPI on a scale of "
        "1 (Very Low Impact) to 5 (Very High Impact). Each barrier has a pre-defined "
        "weight (0, 1, 2, or 3) for every KPI. The higher the weight, the more that "
        "KPI is influenced by that specific barrier.")
    doc.add_paragraph()

    # Section 1
    add_heading(doc, "Section 1: Input Variables (KPI Factors)", level=1)
    add_body(doc, "Each KPI factor below is rated by the user on a scale of 1 to 5.")
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    header_row(tbl, ["#", "KPI Factor Name", "Variable Name"])
    for num, name, var in KPI_FACTORS:
        row = tbl.add_row()
        data_cell(row.cells[0], num, align=WD_ALIGN_PARAGRAPH.CENTER)
        data_cell(row.cells[1], name)
        data_cell(row.cells[2], var, italic=True)
    doc.add_paragraph()

    # Section 2
    add_heading(doc, "Section 2: Score Formula", level=1)
    add_body(doc,
        "For each barrier i (where i = 1 to 15), the KPI Factor Score is calculated as "
        "the weighted sum of all 17 KPI factor input values:")
    doc.add_paragraph()

    add_formula_para(doc,
        "KPI_Factor_Score(Barrier_i)  =  \u03A3 [KPI_j  \u00D7  W(i, j)]   for j = 1 to 17")

    doc.add_paragraph()
    add_body(doc, "Where:")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("KPI_j").italic = True
    p.add_run("  = Input value (1\u20135) of KPI Factor j").font.size = Pt(10)

    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("W(i, j)").italic = True
    p2.add_run("  = Weight of KPI Factor j for Barrier i  (values: 0, 1, 2, or 3)").font.size = Pt(10)

    doc.add_paragraph()
    add_body(doc,
        "Weight Meaning:  0 = No relevance  |  1 = Low relevance  |  "
        "2 = Medium relevance  |  3 = High relevance")
    doc.add_paragraph()

    # Section 3 – Weight matrix
    add_heading(doc, "Section 3: Barrier \u00D7 KPI Factor Weight Matrix", level=1)
    add_body(doc,
        "The table below shows the weight W(i, j) assigned to each KPI factor "
        "for each barrier.")
    doc.add_paragraph()

    num_kpi = len(KPI_FACTORS)
    tbl2 = doc.add_table(rows=1, cols=num_kpi + 2)
    tbl2.style = 'Table Grid'

    hrow = tbl2.rows[0]
    data_cell(hrow.cells[0], "B#", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg='1F4E79')
    hrow.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    data_cell(hrow.cells[1], "Barrier Title", bold=True, bg='1F4E79')
    hrow.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for j, (num, name, var) in enumerate(KPI_FACTORS):
        cell = hrow.cells[j + 2]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(num)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, '1F4E79')
        set_cell_borders(cell)

    W_COLORS = {0: None, 1: 'D6E4F0', 2: '2E74B5', 3: '1F4E79'}
    W_TEXT_COLORS = {0: '000000', 1: '000000', 2: 'FFFFFF', 3: 'FFFFFF'}

    for i, (title, weights) in enumerate(zip(BARRIER_TITLES, KPI_WEIGHTS)):
        row = tbl2.add_row()
        data_cell(row.cells[0], str(i + 1), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        data_cell(row.cells[1], title, size=8)
        for j, w in enumerate(weights):
            cell = row.cells[j + 2]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(w) if w > 0 else '')
            run.font.size = Pt(8)
            run.bold = (w == 3)
            bg = W_COLORS[w]
            if bg:
                set_cell_bg(cell, bg)
                run.font.color.rgb = RGBColor(
                    int(W_TEXT_COLORS[w][0:2], 16),
                    int(W_TEXT_COLORS[w][2:4], 16),
                    int(W_TEXT_COLORS[w][4:6], 16))
            set_cell_borders(cell)

    doc.add_paragraph()

    # Section 4 – Detailed table
    add_heading(doc, "Section 4: Per-Barrier Detailed Formula Table", level=1)
    add_body(doc,
        "The table below lists every KPI factor with its weight for each barrier.")
    doc.add_paragraph()

    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = 'Table Grid'
    header_row(tbl3, ["Barrier\nNumber", "Barrier Title", "KPI Factor", "Weight W(i,j)"])

    ROW_BG = ['F2F7FC', 'FFFFFF']
    for i, (title, weights) in enumerate(zip(BARRIER_TITLES, KPI_WEIGHTS)):
        for j, w in enumerate(weights):
            row = tbl3.add_row()
            bg = ROW_BG[i % 2]
            data_cell(row.cells[0], str(i + 1), align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
            data_cell(row.cells[1], title, size=8, bg=bg)
            data_cell(row.cells[2], KPI_FACTORS[j][1], size=8, bg=bg)
            cell_bg = bg if w < 3 else 'D6E4F0'
            data_cell(row.cells[3], str(w), align=WD_ALIGN_PARAGRAPH.CENTER,
                      bold=(w == 3), bg=cell_bg)

    path = f"{OUTPUT_DIR}\\KPI Factor Score Calc Formula.docx"
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# Document 3: Impact Value Score Calc Formula
# ─────────────────────────────────────────────────────────────────────────────

def make_impact_value_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    add_heading(doc, "Impact Value Score Calculation Formula", level=0)
    doc.add_paragraph()

    add_heading(doc, "Overview", level=1)
    add_body(doc,
        "The Impact Value (IV) is the final composite score assigned to each of the 15 "
        "IoT adoption barriers. It combines three independently calculated component "
        "scores — the Barrier Score, the Cost Factor Score, and the KPI Factor Score — "
        "through a normalization and weighted aggregation process. The barrier with the "
        "highest Impact Value represents the most critical obstacle to IoT adoption for "
        "the assessed organisation.")
    doc.add_paragraph()

    # Step overview box
    add_heading(doc, "Calculation Steps at a Glance", level=2)
    steps = [
        ("Step 1", "Calculate raw Barrier Score for each barrier (B_i)"),
        ("Step 2", "Calculate raw Cost Factor Score for each barrier (C_i)"),
        ("Step 3", "Calculate raw KPI Factor Score for each barrier (K_i)"),
        ("Step 4", "Normalize each score across all 15 barriers"),
        ("Step 5", "Compute weighted Impact Value using normalised scores"),
    ]
    tbl_steps = doc.add_table(rows=1, cols=2)
    tbl_steps.style = 'Table Grid'
    header_row(tbl_steps, ["Step", "Description"])
    for step, desc in steps:
        row = tbl_steps.add_row()
        data_cell(row.cells[0], step, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg='D6E4F0')
        data_cell(row.cells[1], desc)
    doc.add_paragraph()

    # ── Step 1 ───────────────────────────────────────────────────────────────
    add_heading(doc, "Step 1: Barrier Score  (B_i)", level=1)
    add_body(doc,
        "The Barrier Score for each barrier is derived from the user's responses to "
        "barrier-specific survey questions (see Barrier Score Calc Formula document). "
        "Each barrier score is capped to the range [0, 10].")
    add_formula_para(doc, "B_i  =  Barrier Score for Barrier i   (0 \u2264 B_i \u2264 10)")
    doc.add_paragraph()

    # ── Step 2 ───────────────────────────────────────────────────────────────
    add_heading(doc, "Step 2: Cost Factor Score  (C_i)", level=1)
    add_body(doc,
        "The Cost Factor Score for each barrier is derived from the user's cost factor "
        "ratings and the barrier-specific cost factor weight matrix "
        "(see Cost Factor Score Calc Formula document).")
    add_formula_para(doc,
        "C_i  =  \u03A3 [CF_j  \u00D7  W_cost(i, j)]   for j = 1 to 20")
    doc.add_paragraph()
    add_body(doc, "Where:")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("CF_j").italic = True
    p.add_run("  = User-rated value (1\u20135) of Cost Factor j").font.size = Pt(10)
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("W_cost(i, j)").italic = True
    p2.add_run("  = Weight of Cost Factor j for Barrier i  (0, 1, 2, or 3)").font.size = Pt(10)
    doc.add_paragraph()

    # ── Step 3 ───────────────────────────────────────────────────────────────
    add_heading(doc, "Step 3: KPI Factor Score  (K_i)", level=1)
    add_body(doc,
        "The KPI Factor Score for each barrier is derived from the user's KPI ratings "
        "and the barrier-specific KPI weight matrix "
        "(see KPI Factor Score Calc Formula document).")
    add_formula_para(doc,
        "K_i  =  \u03A3 [KPI_j  \u00D7  W_kpi(i, j)]   for j = 1 to 17")
    doc.add_paragraph()
    add_body(doc, "Where:")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("KPI_j").italic = True
    p.add_run("  = User-rated value (1\u20135) of KPI Factor j").font.size = Pt(10)
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("W_kpi(i, j)").italic = True
    p2.add_run("  = Weight of KPI Factor j for Barrier i  (0, 1, 2, or 3)").font.size = Pt(10)
    doc.add_paragraph()

    # ── Step 4 ───────────────────────────────────────────────────────────────
    add_heading(doc, "Step 4: Normalization", level=1)
    add_body(doc,
        "Each score is normalized against the sum of that score type across all 15 "
        "barriers. This converts raw scores into relative proportions so that all three "
        "component types are on a comparable scale before being combined.")
    doc.add_paragraph()

    add_heading(doc, "Barrier Score Normalization:", level=2)
    add_formula_para(doc,
        "B_norm_i  =  B_i  /  \u03A3(B_1 + B_2 + \u2026 + B_15)")

    doc.add_paragraph()
    add_heading(doc, "Cost Factor Score Normalization:", level=2)
    add_formula_para(doc,
        "C_norm_i  =  C_i  /  \u03A3(C_1 + C_2 + \u2026 + C_15)")

    doc.add_paragraph()
    add_heading(doc, "KPI Factor Score Normalization:", level=2)
    add_formula_para(doc,
        "K_norm_i  =  K_i  /  \u03A3(K_1 + K_2 + \u2026 + K_15)")

    doc.add_paragraph()
    add_body(doc,
        "Note: If the total sum of a score type equals zero (all barriers score 0), "
        "the normalized value for all barriers of that type is set to 0 to avoid "
        "division by zero.")
    doc.add_paragraph()

    # ── Step 5 ───────────────────────────────────────────────────────────────
    add_heading(doc, "Step 5: Impact Value Calculation", level=1)
    add_body(doc,
        "The final Impact Value for each barrier is the weighted sum of the three "
        "normalized component scores:")
    doc.add_paragraph()

    add_formula_para(doc,
        "IV_i  =  (W_B \u00D7 B_norm_i)  +  (W_C \u00D7 C_norm_i)  +  (W_K \u00D7 K_norm_i)")

    doc.add_paragraph()

    # Weight values table
    add_heading(doc, "Component Weights:", level=2)
    tbl_w = doc.add_table(rows=1, cols=4)
    tbl_w.style = 'Table Grid'
    header_row(tbl_w, ["Component", "Symbol", "Weight Value", "Rationale"])
    rows_data = [
        ("Barrier Score",      "W_B", f"{WEIGHT_BARRIER:.1f}",
         "Reflects the direct severity of the barrier as measured by indicator responses"),
        ("Cost Factor Score",  "W_C", f"{WEIGHT_COST:.1f}",
         "Reflects the financial burden associated with each barrier"),
        ("KPI Factor Score",   "W_K", f"{WEIGHT_KPI:.1f}",
         "Reflects the operational performance impact of each barrier (highest weight)"),
    ]
    for comp, sym, val, rat in rows_data:
        row = tbl_w.add_row()
        data_cell(row.cells[0], comp, bold=True)
        data_cell(row.cells[1], sym,  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        data_cell(row.cells[2], val,  bold=True,   align=WD_ALIGN_PARAGRAPH.CENTER)
        data_cell(row.cells[3], rat,  size=9)
    doc.add_paragraph()

    add_body(doc,
        f"Note: W_B + W_C + W_K = {WEIGHT_BARRIER} + {WEIGHT_COST} + {WEIGHT_KPI} = "
        f"{WEIGHT_BARRIER + WEIGHT_COST + WEIGHT_KPI:.1f}  (weights sum to 1.0)")
    doc.add_paragraph()

    # ── Full expanded formula ─────────────────────────────────────────────────
    add_heading(doc, "Complete Expanded Formula", level=1)
    add_body(doc,
        "Substituting the weight values into the Impact Value formula gives:")
    doc.add_paragraph()

    add_formula_para(doc,
        f"IV_i  =  ({WEIGHT_BARRIER} \u00D7 B_norm_i)  +  "
        f"({WEIGHT_COST} \u00D7 C_norm_i)  +  "
        f"({WEIGHT_KPI} \u00D7 K_norm_i)")

    doc.add_paragraph()

    # ── Symbol summary ────────────────────────────────────────────────────────
    add_heading(doc, "Symbol Summary", level=1)
    tbl_sym = doc.add_table(rows=1, cols=3)
    tbl_sym.style = 'Table Grid'
    header_row(tbl_sym, ["Symbol", "Meaning", "Source"])
    symbols = [
        ("i",        "Barrier index (1 to 15)",                         "Fixed"),
        ("j",        "Factor index (1\u201320 for cost; 1\u201317 for KPI)", "Fixed"),
        ("B_i",      "Raw Barrier Score for barrier i",                 "Barrier Score Calc"),
        ("C_i",      "Raw Cost Factor Score for barrier i",             "Cost Factor Score Calc"),
        ("K_i",      "Raw KPI Factor Score for barrier i",              "KPI Factor Score Calc"),
        ("B_norm_i", "Normalised Barrier Score for barrier i",          "Step 4"),
        ("C_norm_i", "Normalised Cost Factor Score for barrier i",      "Step 4"),
        ("K_norm_i", "Normalised KPI Factor Score for barrier i",       "Step 4"),
        ("W_B",      f"Weight for Barrier Score = {WEIGHT_BARRIER}",    "settings.py"),
        ("W_C",      f"Weight for Cost Factor Score = {WEIGHT_COST}",   "settings.py"),
        ("W_K",      f"Weight for KPI Factor Score = {WEIGHT_KPI}",     "settings.py"),
        ("IV_i",     "Final Impact Value for barrier i",                "Output"),
    ]
    ROW_BG = ['F2F7FC', 'FFFFFF']
    for idx, (sym, meaning, source) in enumerate(symbols):
        row = tbl_sym.add_row()
        bg = ROW_BG[idx % 2]
        data_cell(row.cells[0], sym,     bold=True,  italic=True, bg=bg)
        data_cell(row.cells[1], meaning, size=9, bg=bg)
        data_cell(row.cells[2], source,  size=9, italic=True, bg=bg)
    doc.add_paragraph()

    # ── Output ────────────────────────────────────────────────────────────────
    add_heading(doc, "Output", level=1)
    add_body(doc,
        "The calculation produces one Impact Value (IV_i) for each of the 15 barriers. "
        "The barriers are then ranked from highest to lowest IV. The top 3 barriers by "
        "Impact Value are highlighted as the most critical areas requiring attention for "
        "successful IoT adoption.")
    doc.add_paragraph()

    tbl_out = doc.add_table(rows=1, cols=5)
    tbl_out.style = 'Table Grid'
    header_row(tbl_out, ["Barrier #", "Barrier Title",
                          "Barrier Score (B_i)", "Cost Score (C_i)", "KPI Score (K_i)"])
    for i, title in enumerate(BARRIER_TITLES):
        row = tbl_out.add_row()
        bg = ROW_BG[i % 2]
        data_cell(row.cells[0], str(i + 1), align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        data_cell(row.cells[1], title, size=9, bg=bg)
        data_cell(row.cells[2], "B_" + str(i + 1), italic=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        data_cell(row.cells[3], "C_" + str(i + 1), italic=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        data_cell(row.cells[4], "K_" + str(i + 1), italic=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)

    path = f"{OUTPUT_DIR}\\Impact Value Score Calc Formula.docx"
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    make_cost_factor_doc()
    make_kpi_factor_doc()
    make_impact_value_doc()
    print("\nAll three documents generated successfully.")
